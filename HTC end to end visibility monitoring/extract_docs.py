import subprocess
import sys

# Install required packages
subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx", "openpyxl", "--quiet"])

from docx import Document
import openpyxl
import json

base = r"c:\Users\solomonal\OneDrive - Ethiopian Airlines\Others\11 KPI folder\2. Projects\10. Hard time missing mandatory tracker"

# Extract DOCX
doc = Document(f"{base}\\HTC_EndToEnd_Visibility_Monitoring_Proposal.docx")
with open(f"{base}\\proposal_text.txt", "w", encoding="utf-8") as f:
    for para in doc.paragraphs:
        f.write(para.text + "\n")
    # Also extract tables
    for i, table in enumerate(doc.tables):
        f.write(f"\n--- TABLE {i+1} ---\n")
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            f.write(" | ".join(cells) + "\n")

print("DOCX extracted successfully")

# Extract XLSX
wb = openpyxl.load_workbook(f"{base}\\JUN-17-2026.xlsx", data_only=True)
with open(f"{base}\\excel_data.txt", "w", encoding="utf-8") as f:
    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        f.write(f"\n=== SHEET: {sheet_name} ===\n")
        for row in ws.iter_rows(values_only=False):
            cells = []
            for cell in row:
                val = cell.value if cell.value is not None else ""
                cells.append(str(val))
            f.write(" | ".join(cells) + "\n")

print("XLSX extracted successfully")
